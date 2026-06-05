from __future__ import annotations

import sys
from pathlib import Path
from docx import Document
from db import get_conn
from llm import embed_text

MAX_CHARS = 1200
OVERLAP = 150


def read_docx(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # 표 안의 텍스트도 함께 수집합니다.
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def chunk_text(text: str) -> list[str]:
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + MAX_CHARS, len(text))
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end == len(text):
            break
        start = max(0, end - OVERLAP)
    return chunks


def main() -> None:
    if len(sys.argv) < 2:
        raise SystemExit('사용법: python ingest_docx.py "기획초안.docx"')

    path = Path(sys.argv[1])
    if not path.exists():
        raise FileNotFoundError(path)

    text = read_docx(path)
    chunks = chunk_text(text)

    with get_conn() as conn:
        doc_id = conn.execute(
            "INSERT INTO documents(title, source_path) VALUES (%s, %s) RETURNING id",
            (path.stem, str(path)),
        ).fetchone()[0]

        for i, chunk in enumerate(chunks, start=1):
            emb = embed_text(chunk)
            conn.execute(
                """
                INSERT INTO chunks(document_id, section_title, content, embedding)
                VALUES (%s, %s, %s, %s)
                """,
                (doc_id, f"chunk-{i}", chunk, emb),
            )

    print(f"ingested document_id={doc_id}, chunks={len(chunks)}")

if __name__ == "__main__":
    main()
